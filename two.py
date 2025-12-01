import os
import re
import vk_api
from telebot import TeleBot, types
from vk_api.exceptions import ApiError
from dotenv import load_dotenv
import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# === Загрузка переменных окружения ===
load_dotenv()
TELEGRAM_TOKEN = os.getenv('TELEGRAM_TOKEN')
VK_TOKEN = os.getenv('VK_TOKEN')
YOUR_CHAT_ID = os.getenv('YOUR_CHAT_ID')

if not TELEGRAM_TOKEN or not VK_TOKEN:
    print("Ошибка: не найдены токены в .env!")
    exit()

bot = TeleBot(TELEGRAM_TOKEN)
vk_session = vk_api.VkApi(token=VK_TOKEN)
vk = vk_session.get_api()

user_states = {}

# === ФУНКЦИИ СОЗДАНИЯ ОТЧЁТОВ ===

def create_doxc_report(groups_data, filename="vk_analysis_report.docx"):
    doc = Document()
    title = doc.add_heading('Анализ групп Вконтакте', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for i, group in enumerate(groups_data, 1):
        doc.add_heading(f"{i}. {group.get('name', 'Без названия')}", level=2)
        
        p = doc.add_paragraph()
        p.add_run("Участников: ").bold = True
        p.add_run(str(group.get("members", "—")))
        
        p = doc.add_paragraph()
        p.add_run("Среднее количество лайков: ").bold = True
        p.add_run(str(group.get("avg_likes", "—")))
        
        if group.get("description"):
            p = doc.add_paragraph()
            p.add_run("Описание: ").bold = True
            p.add_run(group["description"])
        
        doc.add_paragraph()  # отступ

    doc.save(filename)
    return filename


def create_activity_docx(user_info, posts_data, filename_prefix="Активность"):
    """
    Создаёт .docx отчёт по активности пользователя в группе
    """
    doc = Document()
    title = doc.add_heading('Анализ активности пользователя', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Инфо о пользователе и группе
    p = doc.add_paragraph()
    p.add_run("👤 Пользователь: ").bold = True
    p.add_run(user_info["name"])

    p = doc.add_paragraph()
    p.add_run("🔗 Профиль: ").bold = True
    p.add_run(user_info["link"])

    p = doc.add_paragraph()
    p.add_run("👥 Группа: ").bold = True
    p.add_run(user_info["group_name"])

    doc.add_paragraph()

    # Статистика
    total = len(posts_data)
    liked = sum(1 for p in posts_data if p["liked"])
    reposted = sum(1 for p in posts_data if p["reposted"])
    
    doc.add_heading("📊 Сводка", level=2)
    doc.add_paragraph(f"Всего постов: {total}")
    doc.add_paragraph(f"Лайков: {liked}")
    doc.add_paragraph(f"Репостов: {reposted}")
    if total > 0:
        activity_percent = (liked + reposted) / total * 100
        doc.add_paragraph(f"Активность: {activity_percent:.1f}%")
    else:
        doc.add_paragraph("Активность: 0%")
    doc.add_paragraph()

    # Подробности по постам
    doc.add_heading("📝 Детали по постам", level=2)
    for i, item in enumerate(posts_data, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {item['date']}").bold = True
        p.add_run(f"\nСсылка: {item['link']}")
        p.add_run(f"\nЛайк: {'✅ Да' if item['liked'] else '❌ Нет'}")
        p.add_run(f"\nРепост: {'✅ Да' if item['reposted'] else '❌ Нет'}")
        doc.add_paragraph()

    filename = f"{filename_prefix}_{datetime.datetime.now().strftime('%d-%m-%Y_%H-%M')}.docx"
    doc.save(filename)
    return filename


def create_likers_docx(post_info, likers_data, filename_prefix="Лайкнувшие"):
    """
    Создаёт .docx отчёт по лайкнувшим пост
    """
    doc = Document()
    title = doc.add_heading('Список лайкнувших пост', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_paragraph("🔗 Ссылка на пост: ").add_run(post_info["link"]).underline = True
    doc.add_paragraph(f"👥 Всего лайков: {len(likers_data)}")
    doc.add_paragraph()

    doc.add_heading("📝 Список пользователей", level=2)
    for i, user in enumerate(likers_data, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(user["name"])
        p.add_run(f" — {user['link']}")

    filename = f"{filename_prefix}_{datetime.datetime.now().strftime('%d-%m-%Y_%H-%M')}.docx"
    doc.save(filename)
    return filename


def create_txt_report(groups_data, filename="vk_analysis_report.txt"):
    lines = []
    lines.append("АНАЛИЗ ГРУПП ВКОНТАКТЕ")
    lines.append("=" * 40)
    lines.append("")

    for i, group in enumerate(groups_data, 1):
        name = group.get('name', 'Без названия')
        members = group.get('members', '—')
        avg_likes = group.get('avg_likes', '—')
        desc = group.get('description', '').strip()

        lines.append(f"{i}. {name}")
        lines.append(f"   Участников: {members}")
        lines.append(f"   Ср. лайков: {avg_likes}")
        if desc:
            lines.append(f"   Описание: {desc}")
        lines.append("")  # пустая строка между группами

    with open(filename, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    return filename


# === Вспомогательные функции ===
def resolve_vk_id(screen_name):
    try:
        resolved = vk.utils.resolveScreenName(screen_name=screen_name.strip())
        if resolved and resolved.get('object_id'):
            return -resolved['object_id'] if resolved['type'] == 'group' else resolved['object_id']
        return None
    except:
        return None


def extract_screen_name(url):
    match = re.search(r'vk\.com/([a-zA-Z0-9._-]+)', url or "")
    return match.group(1) if match else None


def parse_post_link(link):
    patterns = [
        r'vk\.com/wall(-?\d+)_(\d+)',
        r'vk\.com/wall([a-zA-Z0-9._-]+)\?w=wall(-?\d+)_(\d+)',
        r'm\.vk\.com/wall(-?\d+)_(\d+)'
    ]
    for pattern in patterns:
        match = re.search(pattern, link)
        if match:
            if len(match.groups()) == 2:
                return int(match.group(1)), int(match.group(2))
            elif len(match.groups()) == 3:
                return int(match.group(2)), int(match.group(3))
    return None, None


# === ФУНКЦИЯ ОТПРАВКИ ОТЧЕТА ВЛАДЕЛЬЦУ ===
def send_report_to_owner(chat_id, username, message_text, report_type):
    try:
        if YOUR_CHAT_ID:
            clean_text = re.sub('<[^<]+?>', '', message_text)
            clean_text = clean_text.replace('&nbsp;', ' ').replace('&amp;', '&')
            
            report = f"""📊 ОТЧЕТ ОТ БОТА
━━━━━━━━━━━━━━━━━━━━
👤 Пользователь: @{username if username else 'не указан'}
🆔 Chat ID: {chat_id}
📊 Тип отчета: {report_type}
🕒 Время: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}
━━━━━━━━━━━━━━━━━━━━

{clean_text}
━━━━━━━━━━━━━━━━━━━━
✅ Отчет сгенерирован автоматически"""
            
            bot.send_message(YOUR_CHAT_ID, report)
    except Exception as e:
        print(f"Ошибка при отправке отчета владельцу: {e}")


# === Клавиатуры ===
def main_menu_keyboard():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    markup.add("Начать анализ", "Кто лайкнул пост")
    markup.add("Помощь")
    return markup


def cancel_keyboard():
    markup = types.ReplyKeyboardMarkup(resize_keyboard=True)
    markup.add("Отмена")
    return markup


# === КОМАНДА /start ===
@bot.message_handler(commands=['start'])
def start_command(message):
    response = (
        "<b>Привет!</b>\n\n"
        "Я умею:\n"
        "• Проверять, лайкал ли человек посты в группе\n"
        "• Показывать, кто лайкнул любой пост ВК\n\n"
        "Результаты присылаются:\n"
        "• Подробным сообщением\n"
        "• Файлом DOCX\n\n"
        "Выбери функцию:"
    )
    
    bot.send_message(
        message.chat.id,
        response,
        parse_mode="HTML",
        reply_markup=main_menu_keyboard()
    )
    
    username = message.from_user.username if message.from_user.username else "не указан"
    send_report_to_owner(message.chat.id, username, response, "Команда /start")


# === ОСНОВНОЙ ОБРАБОТЧИК ТЕКСТА ===
@bot.message_handler(content_types=['text'])
def handle_text(message):
    chat_id = message.chat.id
    text = message.text.strip()
    username = message.from_user.username if message.from_user.username else "не указан"

    if text == "Начать анализ":
        response = "<b>Отправь ссылку на группу ВК</b>"
        bot.send_message(chat_id, response, parse_mode="HTML", reply_markup=cancel_keyboard())
        user_states[chat_id] = {'step': 'awaiting_group'}
        send_report_to_owner(chat_id, username, response, "Начало анализа")

    elif text == "Кто лайкнул пост":
        response = (
            "Отправь ссылку на любой пост ВК\n"
            "Пример: https://vk.com/wall-123456789_987654"
        )
        bot.send_message(chat_id, response, reply_markup=cancel_keyboard())
        user_states[chat_id] = {'step': 'awaiting_post_link'}
        send_report_to_owner(chat_id, username, response, "Запрос лайкнувших пост")

    elif text == "Помощь":
        help_text = (
            "Доступные команды:\n"
            "• Начать анализ — проверка активности в группе\n"
            "• Кто лайкнул пост — список лайкнувших\n\n"
            "Результаты присылаются:\n"
            "• Подробным сообщением\n"
            "• Файлом DOCX"
        )
        bot.send_message(chat_id, help_text, reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, help_text, "Запрос помощи")

    elif text == "Отмена":
        user_states.pop(chat_id, None)
        bot.send_message(chat_id, "Отменено!", reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, "Пользователь отменил операцию", "Отмена")

    elif user_states.get(chat_id, {}).get('step') == 'awaiting_group':
        screen_name = extract_screen_name(text)
        if not screen_name:
            response = "Не понял ссылку на группу."
            bot.send_message(chat_id, response)
            send_report_to_owner(chat_id, username, response, "Ошибка группы")
            return
        
        group_id = resolve_vk_id(screen_name)
        if not group_id or group_id > 0:
            response = "Это не группа ВК."
            bot.send_message(chat_id, response)
            send_report_to_owner(chat_id, username, response, "Ошибка группы")
            return
        
        user_states[chat_id] = {'step': 'awaiting_user', 'group_id': group_id}
        response = "<b>Группа принята!</b>\n\nТеперь отправь ссылку на профиль человека"
        bot.send_message(chat_id, response, parse_mode="HTML", reply_markup=cancel_keyboard())
        send_report_to_owner(chat_id, username, response, "Группа принята")

    elif user_states.get(chat_id, {}).get('step') == 'awaiting_user':
        screen_name = extract_screen_name(text)
        if not screen_name:
            response = "Не понял ссылку на человека."
            bot.send_message(chat_id, response)
            send_report_to_owner(chat_id, username, response, "Ошибка пользователя")
            return
        
        user_id = resolve_vk_id(screen_name)
        if not user_id or user_id < 0:
            response = "Это не личный профиль."
            bot.send_message(chat_id, response)
            send_report_to_owner(chat_id, username, response, "Ошибка пользователя")
            return

        response = "Анализирую 30 последних постов...\nОжидай 15–30 секунд"
        bot.send_message(chat_id, response, reply_markup=types.ReplyKeyboardRemove())
        send_report_to_owner(chat_id, username, response, "Начало анализа активности")
        analyze_user_activity(chat_id, user_states[chat_id]['group_id'], user_id, username)
        user_states.pop(chat_id, None)

    elif user_states.get(chat_id, {}).get('step') == 'awaiting_post_link':
        if text == "Отмена":
            user_states.pop(chat_id, None)
            bot.send_message(chat_id, "Отменено!", reply_markup=main_menu_keyboard())
            send_report_to_owner(chat_id, username, "Пользователь отменил операцию", "Отмена поста")
            return

        owner_id, post_id = parse_post_link(text)
        if not owner_id or not post_id:
            response = "Не могу распознать ссылку на пост.\nПопробуй скопировать ссылку прямо из приложения ВК."
            bot.send_message(chat_id, response)
            send_report_to_owner(chat_id, username, response, "Ошибка ссылки на пост")
            return

        response = "Собираю лайки... (до 1000 человек)"
        bot.send_message(chat_id, response, reply_markup=types.ReplyKeyboardRemove())
        send_report_to_owner(chat_id, username, response, "Начало сбора лайков")
        get_post_likers(chat_id, owner_id, post_id, username)
        user_states.pop(chat_id, None)

    else:
        response = "Выбери действие:"
        bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, response, "Неизвестная команда")


# === ФУНКЦИЯ: Кто лайкнул пост ===
def get_post_likers(chat_id, owner_id, post_id, username):
    try:
        likes = vk.likes.getList(
            type='post',
            owner_id=owner_id,
            item_id=post_id,
            count=1000,
            extended=1,
            fields='id,first_name,last_name'
        )
        users = likes.get('items', [])

        if not users:
            response = "Никто не лайкнул этот пост"
            bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
            send_report_to_owner(chat_id, username, response, "Результат: нет лайков")
            return

        count = len(users)
        
        # Подробный отчёт сообщением
        link_clean = f"https://vk.com/wall{owner_id}_{post_id}"
        report = f"<b>📊 Лайкнули пост: {count} человек</b>\n\n"
        report += f"<b>Ссылка на пост:</b>\n{link_clean}\n\n"
        1
        user_list = []
        likers_data = []

        for i, user in enumerate(users[:50], 1):
            name = f"{user['first_name']} {user['last_name']}"
            link = f"https://vk.com/id{user['id']}"
            user_list.append(f"{i}. <a href='{link}'>{name}</a>")
            likers_data.append({"name": name, "link": link})

        report += "<b>Список лайкнувших:</b>\n" + "\n".join(user_list)
        if count > 50:
            report += f"\n\n...и еще {count - 50} человек"
        
        bot.send_message(chat_id, report, parse_mode="HTML", disable_web_page_preview=True)

        # Создаём и отправляем DOCX
        post_info = {"link": link_clean}
        docx_path = create_likers_docx(post_info, likers_data)

        with open(docx_path, 'rb') as f:
            bot.send_document(chat_id, f, caption=f"📎 Список лайкнувших ({count} чел.)")

        os.remove(docx_path)

        bot.send_message(chat_id, "✅ Готово! Все данные отправлены.", reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, report, f"Результат лайков поста ({count} человек)")

    except ApiError as e:
        if e.code == 15:
            response = "❌ Лайки скрыты у этого поста"
            bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
            send_report_to_owner(chat_id, username, response, "Ошибка: лайки скрыты")
        else:
            response = f"❌ Ошибка ВК: {e}"
            bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
            send_report_to_owner(chat_id, username, response, "Ошибка ВК API")
    except Exception as e:
        print("Ошибка:", e)
        response = "❌ Произошла ошибка при получении лайков"
        bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, response, "Ошибка получения лайков")


# === ФУНКЦИЯ: Анализ активности пользователя в группе ===
def analyze_user_activity(chat_id, group_id, user_id, username):
    try:
        posts = vk.wall.get(owner_id=group_id, count=30)['items']
        if not posts:
            response = "❌ Нет постов или доступ закрыт."
            bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
            send_report_to_owner(chat_id, username, response, "Ошибка: нет постов")
            return

        # Получим данные пользователя и группы для заголовка
        try:
            user_info_vk = vk.users.get(user_ids=user_id, fields="first_name,last_name")[0]
            user_name = f"{user_info_vk['first_name']} {user_info_vk['last_name']}"
            user_link = f"https://vk.com/id{user_id}"
        except:
            user_name = "Пользователь"
            user_link = "—"

        try:
            group_info = vk.groups.getById(group_id=-group_id)[0]
            group_name = group_info["name"]
        except:
            group_name = "Группа"

        user_info = {"name": user_name, "link": user_link, "group_name": group_name}

        # Собираем данные по постам
        posts_data = []
        liked = []
        reposted = []
        total_likes = total_reposts = 0

        for post in posts:
            post_id = post['id']
            date_str = datetime.datetime.fromtimestamp(post['date']).strftime("%d.%m.%Y %H:%M")
            link = f"https://vk.com/wall{group_id}_{post_id}"

            try:
                info = vk.likes.isLiked(user_id=user_id, type='post', owner_id=group_id, item_id=post_id)
                has_like = bool(info.get('liked', False))
                has_repost = bool(info.get('copied', False))
            except:
                has_like = has_repost = False

            if has_like:
                total_likes += 1
                liked.append(f"• Пост от {date_str} ({link})")
            if has_repost:
                total_reposts += 1
                reposted.append(f"• Пост от {date_str} ({link})")

            posts_data.append({
                "date": date_str,
                "link": link,
                "liked": has_like,
                "reposted": has_repost
            })

        # Текстовый отчёт
        report = "<b>📊 Анализ завершён!</b>\n\n"
        report += f"<b>Статистика:</b>\n"
        report += f"• Проверено постов: <b>{len(posts)}</b>\n"
        report += f"• Лайков: <b>{total_likes}</b>\n"
        report += f"• Репостов: <b>{total_reposts}</b>\n"
        report += f"• Всего активности: <b>{total_likes + total_reposts}</b>\n"

        if posts:
            activity_percent = (total_likes + total_reposts) / len(posts) * 100
            report += f"• Процент активности: <b>{activity_percent:.1f}%</b>\n\n"
        else:
            report += "\n"

        if liked:
            report += f"<b>❤️ Лайкнутые посты ({total_likes}):</b>\n"
            for item in liked[:10]:
                report += f"{item}\n"
            if len(liked) > 10:
                report += f"...и еще {len(liked) - 10} постов\n\n"
            else:
                report += "\n"

        if reposted:
            report += f"<b>🔄 Репосты ({total_reposts}):</b>\n"
            for item in reposted[:10]:
                report += f"{item}\n"
            if len(reposted) > 10:
                report += f"...и еще {len(reposted) - 10} постов\n"

        if not liked and not reposted:
            report += "😴 Пользователь <b>ничего не лайкал и не репостил</b>."

        bot.send_message(chat_id, report, parse_mode="HTML", disable_web_page_preview=True)

        # Генерация DOCX
        docx_path = create_activity_docx(user_info, posts_data)

        with open(docx_path, 'rb') as f:
            bot.send_document(chat_id, f, caption="📎 Подробный отчёт в формате DOCX")

        os.remove(docx_path)

        bot.send_message(chat_id, "✅ Готово! Все данные отправлены.", reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, report, f"Результат анализа активности")

    except Exception as e:
        print("Ошибка анализа:", e)
        response = "❌ Ошибка при анализе."
        bot.send_message(chat_id, response, reply_markup=main_menu_keyboard())
        send_report_to_owner(chat_id, username, response, "Ошибка анализа")


# === Запуск ===
if __name__ == '__main__':
    print("✅ Бот запущен — отправляет результаты сообщениями и DOCX-файлами!")
    bot.polling(none_stop=True, interval=0)